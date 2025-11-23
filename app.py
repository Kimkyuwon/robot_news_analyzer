import streamlit as st
import google.generativeai as genai
from duckduckgo_search import DDGS
from datetime import datetime, timedelta
import time
import os
import PyPDF2
import io
import json
from docx import Document
from docx.oxml.ns import qn
from fpdf import FPDF

# Page configuration
st.set_page_config(
    page_title="로봇 산업 주간 분석 리포트",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for mobile responsiveness and better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        margin-bottom: 1rem;
    }
    .section-header {
        font-size: 1.8rem;
        font-weight: bold;
        color: #ff7f0e;
        margin-top: 2rem;
        margin-bottom: 1rem;
        border-bottom: 3px solid #ff7f0e;
        padding-bottom: 0.5rem;
    }
    .news-card {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
        border-left: 4px solid #1f77b4;
    }
    .news-title {
        font-weight: bold;
        color: #2c3e50;
        margin-bottom: 0.5rem;
    }
    .news-snippet {
        color: #555;
        font-size: 0.9rem;
        margin-bottom: 0.5rem;
    }
    .news-link {
        font-size: 0.85rem;
        color: #1f77b4;
    }
    @media (max-width: 768px) {
        .main-header {
            font-size: 1.8rem;
        }
        .section-header {
            font-size: 1.4rem;
        }
    }
</style>
""", unsafe_allow_html=True)

# API Key file path
API_KEY_FILE = os.path.join(os.path.dirname(__file__), '.api_key.txt')
HISTORY_FILE = os.path.join(os.path.dirname(__file__), '.analysis_history.json')

# Function to load API key from file
def load_api_key():
    if os.path.exists(API_KEY_FILE):
        try:
            with open(API_KEY_FILE, 'r') as f:
                return f.read().strip()
        except:
            return ""
    return ""

# Function to save API key to file
def save_api_key(api_key):
    try:
        with open(API_KEY_FILE, 'w') as f:
            f.write(api_key)
        return True
    except:
        return False

# Function to load analysis history
def load_history():
    if os.path.exists(HISTORY_FILE):
        try:
            with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []

# Function to save analysis to history
def save_to_history(analysis_type, content):
    try:
        history = load_history()
        
        # Keep only last 10 analyses
        if len(history) >= 10:
            history = history[-9:]
        
        history.append({
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'type': analysis_type,
            'content': content[:1000]  # Save first 1000 chars as summary
        })
        
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            json.dump(history, f, ensure_ascii=False, indent=2)
        
        return True
    except Exception as e:
        st.warning(f"히스토리 저장 실패: {str(e)}")
        return False

        return False

# Function to delete history item
def delete_history_item(index):
    try:
        history = load_history()
        if 0 <= index < len(history):
            del history[index]
            with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
            return True
        return False
    except:
        return False

# Function to get history summary
def get_history_summary(selected_indices=None):
    history = load_history()
    if not history:
        return "이전 분석 기록이 없습니다."
    
    summary = "=== 이전 분석 히스토리 ===\n\n"
    
    # Filter by selected indices if provided
    if selected_indices is not None:
        target_history = [history[i] for i in selected_indices if 0 <= i < len(history)]
    else:
        target_history = history[-5:]  # Default to last 5
        
    if not target_history:
        return "선택된 이전 분석 기록이 없습니다."
        
    for i, item in enumerate(target_history, 1):
        summary += f"{i}. [{item['timestamp']}] {item['type']}\n"
        summary += f"   요약: {item['content'][:200]}...\n\n"
    
    return summary

# Function to save as Word
def save_to_word(content):
    try:
        doc = Document()
        
        # Set style for Korean font
        style = doc.styles['Normal']
        style.font.name = 'Malgun Gothic'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        
        # Add heading
        heading = doc.add_heading('로봇 산업 분석 리포트', 0)
        heading.style.font.name = 'Malgun Gothic'
        heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        
        # Add timestamp
        p = doc.add_paragraph(f"생성 일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        p.style = doc.styles['Normal']
        
        doc.add_paragraph("-" * 50)
        
        # Add content
        for line in content.split('\n'):
            p = doc.add_paragraph(line)
            p.style = doc.styles['Normal']
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Word 생성 실패: {str(e)}")
        return None

# Function to save as PDF
def save_to_pdf(content):
    try:
        pdf = FPDF()
        pdf.add_page()
        
        # Add a Unicode font (using Malgun Gothic for Korean support)
        # Check for Windows font first, then try Linux fallback
        font_path = "C:/Windows/Fonts/malgun.ttf"
        if not os.path.exists(font_path):
            # Try common Linux Korean fonts or fallback
            possible_paths = [
                "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
                "/usr/share/fonts/nanum/NanumGothic.ttf"
            ]
            for path in possible_paths:
                if os.path.exists(path):
                    font_path = path
                    break
        
        if os.path.exists(font_path):
            pdf.add_font('Korean', '', font_path, uni=True)
            pdf.set_font('Korean', '', 11)
        else:
            st.warning("한글 폰트를 찾을 수 없어 기본 폰트를 사용합니다. 한글이 깨질 수 있습니다.")
            pdf.set_font("Arial", size=11)
            
        pdf.cell(0, 10, "로봇 산업 분석 리포트", new_x="LMARGIN", new_y="NEXT", align='C')
        pdf.cell(0, 10, f"생성 일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", new_x="LMARGIN", new_y="NEXT", align='R')
        pdf.ln(10)
        
        # Split content by lines and write
        # Replace unsupported characters
        content = content.replace('\u2022', '-').replace('\u2013', '-').replace('\u2014', '-')
        
        for line in content.split('\n'):
            # Handle empty lines
            if not line.strip():
                pdf.ln(5)
                continue
            
            # Use multi_cell for automatic wrapping
            try:
                pdf.multi_cell(0, 8, line)
            except Exception:
                # Fallback for problematic lines (e.g. very long words)
                try:
                    pdf.multi_cell(0, 8, line[:100] + "...")
                except:
                    pass
            
        # Output to bytes
        return pdf.output(dest='S').encode('latin-1')
    except Exception as e:
        st.error(f"PDF 생성 실패: {str(e)}")
        return None

# Initialize session state
if 'search_results' not in st.session_state:
    st.session_state.search_results = None
if 'ai_report' not in st.session_state:
    st.session_state.ai_report = None
if 'gemini_api_key' not in st.session_state:
    st.session_state.gemini_api_key = load_api_key()

# Sidebar configuration
with st.sidebar:
    st.markdown("### 🔑 API 설정")
    api_key = st.text_input(
        "Gemini API Key", 
        value=st.session_state.gemini_api_key,
        type="password", 
        help="Google AI Studio에서 발급받은 API 키를 입력하세요"
    )
    
    # Save API key to session state and file
    if api_key and api_key != st.session_state.gemini_api_key:
        st.session_state.gemini_api_key = api_key
        if save_api_key(api_key):
            st.success("✅ API 키가 저장되었습니다!")
        else:
            st.warning("⚠️ API 키 저장 실패")
    
    st.markdown("---")
    st.markdown("### ⚙️ 검색 키워드 설정")
    
    st.markdown("**그룹 A (핵심 - 70%)**")
    group_a_construction = st.text_area(
        "건설 로봇 키워드",
        value="건설 로봇\n건설 현장 자동화\n스마트 건설 R&D\n건설용 웨어러블 로봇",
        height=100
    )
    
    group_a_humanoid = st.text_area(
        "휴머노이드 키워드",
        value="휴머노이드 로봇\n이족보행 로봇\n테슬라 옵티머스\n피규어 AI\n보스턴 다이내믹스",
        height=120
    )
    
    st.markdown("**그룹 B (일반 - 30%)**")
    group_b_keywords = st.text_area(
        "기타 로봇 키워드",
        value="협동로봇\n물류 로봇\nAMR\n주차 로봇\n제조업 로봇",
        height=100
    )
    
    st.markdown("---")
    
    # History option
    st.markdown("### 📚 분석 히스토리")
    use_history = st.checkbox(
        "이전 분석 결과 참고",
        value=True,
        help="체크하면 이전 분석 결과를 참고하여 더 깊이 있는 분석을 제공합니다"
    )
    
    if use_history:
        history = load_history()
        if history:
            st.markdown("##### 🕰️ 히스토리 관리")
            st.markdown(f"<small>총 {len(history)}개의 분석 기록</small>", unsafe_allow_html=True)
            
            # Selected indices for context
            selected_history_indices = []
            
            # Iterate through history items (reverse order to show newest first)
            for i in range(len(history) - 1, -1, -1):
                item = history[i]
                with st.expander(f"{item['timestamp']} ({item['type']})"):
                    st.caption(f"요약: {item['content'][:100]}...")
                    
                    # Selection checkbox
                    if st.checkbox("분석에 포함", value=True, key=f"hist_sel_{i}"):
                        selected_history_indices.append(i)
                    
                    # Delete button
                    if st.button("🗑️ 삭제", key=f"hist_del_{i}"):
                        if delete_history_item(i):
                            st.success("삭제됨")
                            time.sleep(0.5)
                            st.rerun()
        else:
            st.info("📊 저장된 분석이 없습니다")
            selected_history_indices = []
    else:
        selected_history_indices = []
    
    st.markdown("---")
    st.markdown("### 📖 사용 방법")
    st.markdown("""
    1. Gemini API 키를 입력하세요
    2. 필요시 검색 키워드를 수정하세요
    3. 각 탭에서 분석 버튼을 클릭하세요
    4. 분석 완료까지 약 1-2분 소요됩니다
    """)

# Function to search news using DuckDuckGo
def search_news(keywords_list, max_results=5):
    """Search news using DuckDuckGo"""
    all_results = []
    seen_urls = set()
    
    ddgs = DDGS()
    
    for keyword in keywords_list:
        try:
            # Search with time filter (past week)
            results = ddgs.text(
                keyword,
                region='kr-kr',
                safesearch='off',
                timelimit='w',  # Past week
                max_results=max_results
            )
            
            for result in results:
                url = result.get('href', '')
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    all_results.append({
                        'title': result.get('title', ''),
                        'snippet': result.get('body', ''),
                        'url': url,
                        'keyword': keyword
                    })
            
            # Small delay to avoid rate limiting
            time.sleep(0.5)
            
        except Exception as e:
            st.warning(f"검색 실패 (키워드: {keyword}): {str(e)}")
    return all_results

# Function to generate AI report using Gemini
def generate_ai_report(group_a_news, group_b_news, api_key, use_history=False, selected_indices=None):
    """Generate analysis report using Gemini AI"""
    try:
        # Configure Gemini API
        genai.configure(api_key=api_key)
        
        # Get history if requested
        history_context = ""
        if use_history:
            history_context = f"\n\n**이전 분석 참고:**\n{get_history_summary(selected_indices)}\n"
        
        # System instruction
        system_instruction = f"""
너는 로봇 산업 전문 애널리스트야. 제공된 뉴스를 단순히 요약하지 말고, 너의 전문적인 분석과 인사이트를 제공해야 해.
{history_context}
**핵심 지침:**
1. 전체 리포트의 **70%**는 '건설 로봇의 현장 적용'과 '휴머노이드의 기술 진척(제어, AI, 하드웨어)'에 집중
2. 두 분야의 융합 가능성(예: 휴머노이드의 건설 현장 투입)을 적극적으로 분석
3. 나머지 30%는 기타 로봇 시장 동향
4. **중요**: 뉴스를 나열하지 말고, 트렌드를 파악하고 너의 분석을 제시해
5. 이전 분석이 있다면, 트렌드 변화와 연속성을 분석해

**리포트 구조:**

## 1. 🏗️ 건설 로봇 & 휴머노이드 심층 분석 (70%)

### 1.1 건설 로봇 현장 적용 분석
- 현재 기술 수준과 실제 적용 사례 분석
- 주요 기술적 과제와 해결 방향
- 시장 성장 가능성 평가

### 1.2 휴머노이드 로봇 기술 진척
- 제어 기술의 최신 동향 (보행, 균형, 조작)
- AI 통합 현황 (비전, 자율성, 학습)
- 하드웨어 혁신 (액추에이터, 센서, 배터리)

### 1.3 융합 시나리오 분석
- 휴머노이드의 건설 현장 투입 가능성
- 기술적 요구사항과 현재 격차
- 예상 타임라인과 선도 기업

### 1.4 주요 기업 및 프로젝트 평가
- 핵심 플레이어 분석 (테슬라, 보스턴다이내믹스, Figure AI 등)
- 투자 동향과 전략적 방향

## 2. 🤖 기타 로봇 산업 동향 (30%)
- 협동로봇, 물류로봇, AMR 등의 주요 트렌드
- 시장 성장 동력과 제약 요인

## 3. 💡 AI 전망 및 투자 인사이트
- **단기 전망 (6개월~1년)**: 예상되는 주요 이벤트와 기술 발표
- **중기 전망 (1~3년)**: 시장 구조 변화와 기술 성숙도
- **장기 전망 (3~5년)**: 산업 패러다임 전환 가능성
- **투자 관점**: 주목해야 할 기업, 기술, 시장 세그먼트
- **리스크 요인**: 기술적/규제적/시장 리스크

**작성 스타일:**
- 전문적이고 분석적인 톤
- 구체적인 수치와 사례 인용
- 명확한 근거를 바탕으로 한 전망
- 불확실성이 있는 부분은 솔직하게 언급
"""
        
        # Prepare news data
        group_a_text = "\n\n".join([
            f"제목: {news['title']}\n내용: {news['snippet']}\n출처: {news['url']}"
            for news in group_a_news
        ])
        
        group_b_text = "\n\n".join([
            f"제목: {news['title']}\n내용: {news['snippet']}\n출처: {news['url']}"
            for news in group_b_news
        ])
        
        # Create full prompt
        full_prompt = f"""{system_instruction}

다음 뉴스 데이터를 바탕으로 주간 로봇 산업 분석 리포트를 작성해주세요.

[그룹 A - 건설 로봇 & 휴머노이드 뉴스 (핵심)]
{group_a_text}

[그룹 B - 기타 로봇 뉴스]
{group_b_text}

현재 날짜: {datetime.now().strftime('%Y년 %m월 %d일')}
분석 기간: 최근 1주일
"""
        
        # Use google-generativeai library (same as stock advisor)
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content(full_prompt)
        
        # Save to history
        if response.text:
            save_to_history("주간 뉴스 분석", response.text)
        
        return response.text
        
    except Exception as e:
        st.error(f"AI 리포트 생성 실패: {str(e)}")
        return None

# Function to extract text from PDF
def extract_pdf_text(pdf_file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.read()))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"PDF 읽기 실패: {str(e)}")
        return None

# Function to analyze uploaded files
def analyze_files(files, api_key):
    """Analyze uploaded files using Gemini AI"""
    try:
        genai.configure(api_key=api_key)
        
        all_text = ""
        for file in files:
            if file.type == "application/pdf":
                text = extract_pdf_text(file)
                if text:
                    all_text += f"\n\n=== {file.name} ===\n{text}"
            elif file.type == "text/plain":
                text = file.read().decode('utf-8')
                all_text += f"\n\n=== {file.name} ===\n{text}"
        
        if not all_text:
            return None
        
        prompt = f"""
다음 문서들을 분석하여 로봇 산업 관점에서 종합 리포트를 작성해주세요.

**분석 요구사항:**
1. 문서의 주요 내용 요약
2. 로봇 산업과의 연관성 분석
3. 기술적 시사점 및 트렌드
4. 비즈니스 및 투자 인사이트
5. 향후 전망 및 권고사항

**문서 내용:**
{all_text}

**작성 스타일:**
- 전문적이고 분석적인 톤
- 구체적인 내용 인용
- 명확한 구조화
- 실용적인 인사이트 제공
"""
        
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content(prompt)
        
        # Save to history
        if response.text:
            save_to_history("파일 분석", response.text)
        
        return response.text
        
    except Exception as e:
        st.error(f"파일 분석 실패: {str(e)}")
        return None

# Function to generate integrated report
def generate_integrated_report(news_report, file_report, api_key):
    """Generate integrated analysis combining news and file analysis"""
    try:
        genai.configure(api_key=api_key)
        
        prompt = f"""
다음 두 가지 분석 결과를 통합하여 종합 리포트를 작성해주세요.

**분석 1: 주간 뉴스 분석 결과**
{news_report}

**분석 2: 파일 분석 결과**
{file_report}

**통합 리포트 작성 요구사항:**

## 1. 🔄 교차 분석 및 시너지
- 뉴스 트렌드와 파일 내용의 연관성 분석
- 상호 보완적인 인사이트 도출
- 일치하는 부분과 차이점 분석

## 2. 🎯 핵심 인사이트 통합
- 두 분석에서 공통으로 나타나는 핵심 트렌드
- 각 분석에서만 나타나는 독특한 인사이트
- 통합적 관점에서의 시장 전망

## 3. 💡 전략적 제언
- 뉴스와 문서 분석을 종합한 실행 가능한 전략
- 단기/중기/장기 관점의 권고사항
- 주목해야 할 기회와 리스크

## 4. 📊 종합 결론
- 로봇 산업의 현재 상황 종합
- 향후 전망 및 예측
- 최종 투자/비즈니스 인사이트

**작성 스타일:**
- 두 분석을 유기적으로 연결
- 구체적인 근거와 예시 제시
- 실용적이고 실행 가능한 제언
- 명확하고 구조화된 형식
"""
        
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content(prompt)
        
        # Save to history
        if response.text:
            save_to_history("통합 분석", response.text)
        
        return response.text
        
    except Exception as e:
        st.error(f"통합 리포트 생성 실패: {str(e)}")
        return None

# Main content with tabs
st.markdown('<div class="main-header">🤖 로봇 산업 분석 플랫폼</div>', unsafe_allow_html=True)

tab1, tab2, tab3 = st.tabs(["📰 주간 뉴스 분석", "📄 파일 업로드 분석", "🔄 통합 분석"])

# Tab 1: Weekly News Analysis
with tab1:
    st.markdown("**건설 로봇**과 **휴머노이드**를 중심으로 한 로봇 산업 심층 분석")
    
    news_analysis_button = st.button("🔍 뉴스 분석 시작", type="primary", key="news_analysis_btn")
    
    # Generate report when button is clicked
    if news_analysis_button:
        if not api_key:
            st.error("⚠️ Gemini API 키를 입력해주세요!")
        else:
            # Parse keywords
            construction_keywords = [k.strip() for k in group_a_construction.split('\n') if k.strip()]
            humanoid_keywords = [k.strip() for k in group_a_humanoid.split('\n') if k.strip()]
            other_keywords = [k.strip() for k in group_b_keywords.split('\n') if k.strip()]
            
            group_a_all = construction_keywords + humanoid_keywords
            
            # Search progress
            with st.spinner('🔍 뉴스 검색 중... (약 30-60초 소요)'):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Search Group A (high priority)
                status_text.text("그룹 A 검색 중 (건설 로봇 & 휴머노이드)...")
                group_a_results = search_news(group_a_all, max_results=5)
                progress_bar.progress(60)
                
                # Search Group B
                status_text.text("그룹 B 검색 중 (기타 로봇)...")
                group_b_results = search_news(other_keywords, max_results=3)
                progress_bar.progress(80)
                
                status_text.text("검색 완료!")
                progress_bar.progress(100)
                time.sleep(0.5)
                progress_bar.empty()
                status_text.empty()
            
            # Store results
            st.session_state.search_results = {
                'group_a': group_a_results,
                'group_b': group_b_results
            }
            
            # Generate AI report
            if group_a_results or group_b_results:
                with st.spinner('🤖 AI 분석 중... (약 30초 소요)'):
                    ai_report = generate_ai_report(
                        group_a_results, 
                        group_b_results, 
                        api_key, 
                        use_history=use_history,
                        selected_indices=selected_history_indices
                    )
                    st.session_state.ai_report = ai_report
                
                st.success(f"✅ 리포트 생성 완료! (그룹 A: {len(group_a_results)}건, 그룹 B: {len(group_b_results)}건)")
            else:
                st.error("검색 결과가 없습니다. 키워드를 변경해보세요.")
    
    # Display AI report
    if st.session_state.ai_report:
        st.markdown("---")
        st.markdown('<div class="section-header">📊 AI 분석 리포트</div>', unsafe_allow_html=True)
        
        with st.container():
            st.markdown(st.session_state.ai_report)
            
            # Export buttons
            st.markdown("### 💾 리포트 저장")
            col1, col2 = st.columns(2)
            with col1:
                docx_data = save_to_word(st.session_state.ai_report)
                if docx_data:
                    st.download_button(
                        label="📄 Word로 저장",
                        data=docx_data,
                        file_name="주간_로봇_산업_분석.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="save_news_word"
                    )
            with col2:
                pdf_data = save_to_pdf(st.session_state.ai_report)
                if pdf_data:
                    st.download_button(
                        label="📑 PDF로 저장",
                        data=pdf_data,
                        file_name="주간_로봇_산업_분석.pdf",
                        mime="application/pdf",
                        key="save_news_pdf"
                    )
        
        # Show source count at bottom
        if st.session_state.search_results:
            results = st.session_state.search_results
            total_sources = len(results.get('group_a', [])) + len(results.get('group_b', []))
            st.info(f"📰 분석에 사용된 뉴스 소스: 총 {total_sources}건 (건설/휴머노이드: {len(results.get('group_a', []))}건, 기타: {len(results.get('group_b', []))}건)")

# Tab 2: File Upload Analysis
with tab2:
    st.markdown("### 📄 파일 업로드 분석")
    st.markdown("PDF 또는 텍스트 파일을 업로드하여 로봇 산업 관점에서 분석합니다.")
    
    # Initialize session state for file analysis
    if 'file_analysis_report' not in st.session_state:
        st.session_state.file_analysis_report = None
    
    uploaded_files = st.file_uploader(
        "파일 선택 (PDF, TXT)",
        type=['pdf', 'txt'],
        accept_multiple_files=True,
        help="여러 파일을 동시에 업로드할 수 있습니다"
    )
    
    analyze_button = st.button("🔍 파일 분석 시작", type="primary", key="analyze_files_btn")
    
    if analyze_button:
        if not api_key:
            st.error("⚠️ Gemini API 키를 입력해주세요!")
        elif not uploaded_files:
            st.error("⚠️ 분석할 파일을 업로드해주세요!")
        else:
            with st.spinner('📄 파일 분석 중... (약 30-60초 소요)'):
                file_report = analyze_files(uploaded_files, api_key)
                st.session_state.file_analysis_report = file_report
            
            if file_report:
                st.success(f"✅ 분석 완료! ({len(uploaded_files)}개 파일)")
    
    # Display file analysis report
    if st.session_state.file_analysis_report:
        st.markdown("---")
        st.markdown('<div class="section-header">📊 파일 분석 리포트</div>', unsafe_allow_html=True)
        
        with st.container():
            st.markdown(st.session_state.file_analysis_report)
            
            # Export buttons
            st.markdown("### 💾 리포트 저장")
            col1, col2 = st.columns(2)
            with col1:
                docx_data = save_to_word(st.session_state.file_analysis_report)
                if docx_data:
                    st.download_button(
                        label="📄 Word로 저장",
                        data=docx_data,
                        file_name="파일_분석_리포트.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="save_file_word"
                    )
            with col2:
                pdf_data = save_to_pdf(st.session_state.file_analysis_report)
                if pdf_data:
                    st.download_button(
                        label="📑 PDF로 저장",
                        data=pdf_data,
                        file_name="파일_분석_리포트.pdf",
                        mime="application/pdf",
                        key="save_file_pdf"
                    )
        
        if uploaded_files:
            st.info(f"📁 분석된 파일: {', '.join([f.name for f in uploaded_files])}")

# Tab 3: Integrated Analysis
with tab3:
    st.markdown("### 🔄 통합 분석")
    st.markdown("주간 뉴스 분석과 파일 분석 결과를 통합하여 종합적인 인사이트를 제공합니다.")
    
    # Initialize session state for integrated analysis
    if 'integrated_report' not in st.session_state:
        st.session_state.integrated_report = None
    
    # Check if both analyses are available
    has_news = st.session_state.ai_report is not None
    has_files = st.session_state.file_analysis_report is not None
    
    if has_news and has_files:
        st.success("✅ 주간 뉴스 분석과 파일 분석 결과가 모두 준비되었습니다!")
        
        integrate_button = st.button("🔄 통합 분석 시작", type="primary", key="integrate_btn")
        
        if integrate_button:
            if not api_key:
                st.error("⚠️ Gemini API 키를 입력해주세요!")
            else:
                with st.spinner('🔄 통합 분석 중... (약 30-60초 소요)'):
                    integrated_report = generate_integrated_report(
                        st.session_state.ai_report,
                        st.session_state.file_analysis_report,
                        api_key
                    )
                    st.session_state.integrated_report = integrated_report
                
                if integrated_report:
                    st.success("✅ 통합 분석 완료!")
        
        # Display integrated report
        if st.session_state.integrated_report:
            st.markdown("---")
            st.markdown('<div class="section-header">📊 통합 분석 리포트</div>', unsafe_allow_html=True)
            
            with st.container():
                st.markdown(st.session_state.integrated_report)
                
                # Export buttons
                st.markdown("### 💾 리포트 저장")
                col1, col2 = st.columns(2)
                with col1:
                    docx_data = save_to_word(st.session_state.integrated_report)
                    if docx_data:
                        st.download_button(
                            label="📄 Word로 저장",
                            data=docx_data,
                            file_name="통합_분석_리포트.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="save_integrated_word"
                        )
                with col2:
                    pdf_data = save_to_pdf(st.session_state.integrated_report)
                    if pdf_data:
                        st.download_button(
                            label="📑 PDF로 저장",
                            data=pdf_data,
                            file_name="통합_분석_리포트.pdf",
                            mime="application/pdf",
                            key="save_integrated_pdf"
                        )
            
            # Show summary
            st.info("💡 이 리포트는 주간 뉴스 트렌드와 업로드된 문서를 종합적으로 분석한 결과입니다.")
    
    elif has_news and not has_files:
        st.warning("⚠️ 파일 분석 결과가 없습니다. '📄 파일 업로드 분석' 탭에서 파일을 업로드하고 분석해주세요.")
    elif not has_news and has_files:
        st.warning("⚠️ 주간 뉴스 분석 결과가 없습니다. '📰 주간 뉴스 분석' 탭에서 리포트를 생성해주세요.")
    else:
        st.info("ℹ️ 통합 분석을 위해서는 먼저 다음 작업을 완료해주세요:")
        st.markdown("""
        1. 주간 뉴스 분석 탭에서 뉴스 리포트 생성
        2. 파일 업로드 분석 탭에서 파일 분석 완료
        3. 이 탭으로 돌아와서 통합 분석 시작 버튼 클릭
        """)


# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #888; font-size: 0.9rem;">
    <p>Robot Industry Analysis Platform | Powered by Gemini AI & DuckDuckGo</p>
    <p>Generated: {}</p>
</div>
""".format(datetime.now().strftime('%Y-%m-%d %H:%M:%S')), unsafe_allow_html=True)
